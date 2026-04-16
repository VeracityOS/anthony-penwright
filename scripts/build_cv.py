"""Build the Master CV for Anthony Penwright (DOCX + PDF).

Source variants: d:/CV Anthony/*.pdf
Outputs:  d:/Dev/GitHub/anthony-penwright/public/anthony-penwright-cv.{docx,pdf}
"""
from __future__ import annotations
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OXFORD = RGBColor(0x0A, 0x25, 0x40)
SIENNA = RGBColor(0xC8, 0x53, 0x2C)
INK = RGBColor(0x11, 0x11, 0x11)
MUTED = RGBColor(0x55, 0x5B, 0x66)
RULE = RGBColor(0xC9, 0xD1, 0xDB)

OUT_DIR = Path(r"d:/Dev/GitHub/anthony-penwright/public")
OUT_DIR.mkdir(parents=True, exist_ok=True)
DOCX_PATH = OUT_DIR / "anthony-penwright-cv.docx"
PDF_PATH = OUT_DIR / "anthony-penwright-cv.pdf"


def set_margins(doc, top=1.5, bottom=1.5, left=1.8, right=1.8):
    for s in doc.sections:
        s.top_margin = Cm(top)
        s.bottom_margin = Cm(bottom)
        s.left_margin = Cm(left)
        s.right_margin = Cm(right)


def add_run(p, text, *, bold=False, italic=False, size=10, color=INK, font="Calibri",
            smallcaps=False, letter_spacing=None):
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    f = r.font
    f.name = font
    f.size = Pt(size)
    f.color.rgb = color
    if smallcaps:
        f.small_caps = True
    if letter_spacing is not None:
        rPr = r._element.get_or_add_rPr()
        spc = OxmlElement('w:spacing')
        spc.set(qn('w:val'), str(letter_spacing))
        rPr.append(spc)
    return r


def h_rule(doc, color=RULE):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '4')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), f'{color[0]:02X}{color[1]:02X}{color[2]:02X}')
    pbdr.append(bot)
    pPr.append(pbdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)


def section_label(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, text, bold=True, size=9.5, color=OXFORD, font="Calibri",
            smallcaps=True, letter_spacing=40)
    h_rule(doc)
    return p


def subheading(doc, left, right=None, *, size=10.5, space_before=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(1)
    add_run(p, left, bold=True, size=size, color=OXFORD)
    if right:
        add_run(p, "    " + right, bold=False, size=size - 0.5, color=SIENNA, italic=False)
    return p


def meta_line(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, text, size=9, color=MUTED, italic=True)
    return p


def body(doc, text, *, size=10, space_after=3, indent=None, bold=False, color=INK):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    add_run(p, text, size=size, color=color, bold=bold)
    return p


def bullet(doc, text, *, size=10):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    # Clear existing runs
    for r in p.runs:
        r.text = ""
    add_run(p, text, size=size, color=INK)
    return p


def page_break(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r._r.append(br)


def build():
    doc = Document()
    # Base style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    set_margins(doc)

    # ============================================================
    # PAGE 1 — COVER / PROFILE
    # ============================================================
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    add_run(p, "ANTHONY PENWRIGHT", bold=True, size=26, color=OXFORD, font="Georgia")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    add_run(p, "C-Suite Advisor  ·  Smart Cities  ·  Digital Infrastructure  ·  Innovation Ecosystems",
            size=11.5, color=SIENNA, italic=True, font="Georgia")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    add_run(p,
            "ap@credible.uk.com   ·   +966 533 15 8265   ·   Riyadh, KSA (Premium Resident)   ·   "
            "LinkedIn: Anthony Penwright",
            size=9.5, color=MUTED)
    h_rule(doc)

    section_label(doc, "Executive Summary")
    body(doc,
         "C-Suite advisor codifying 20+ years of global smart-city and digital-infrastructure delivery into "
         "frameworks that work. $1.3BN+ of programme value delivered across 40+ major initiatives and 13 sectors "
         "— MoD DII(F) (150,000 workstations · 242 sites · 150+ countries), NEOM ($290M digital master plan · "
         "$120M innovation portfolio), FCO (99.8% SLA across the diplomatic estate), Cisco's global Smart City "
         "practice (built from zero to $315M), and Quantela ($18M → $90M in eighteen months).",
         space_after=5)
    body(doc,
         "Author of four proprietary methodologies — Smart Services Advisory, Innovation Ecosystem Advisory, "
         "Operations Advisory and the Value Realisation Office — live across DMCC Uptown Dubai, NEOM Sindalah, "
         "Red Sea Global and Strive Services Group KSA. Founder & Principal of Verax Venture Studio: AI-native, "
         "7 active portfolio ventures, 44+ agent skills with master-agent orchestration across CRM, KYC, brand "
         "and operations.",
         space_after=5)
    body(doc,
         "Proven architect of high-performance cross-functional teams (450+ resources), C-suite adviser on "
         "IoT/AI/big-data strategy, and experienced navigator of complex regulatory and government approval "
         "environments. Deep expertise across cloud infrastructure (Azure, AWS), OT/IT convergence, ERP "
         "transformation (SAP, Oracle Fusion, Dynamics 365), cybersecurity and data governance — with extensive "
         "GCC/MENA operating experience and a current Riyadh base.",
         space_after=5)
    body(doc,
         "Trusted C-suite advisor who translates complex technical roadmaps into measurable business outcomes. "
         "Dual MBA-qualified (Entrepreneurial Action; Artificial Intelligence in progress). PRINCE2 Practitioner · "
         "MSP · ITIL v3 · CompTIA Security+ · BTA Blockchain Professional. UK SC / DV eligible (FCO & MoD history).",
         space_after=8)

    # Stats strip
    section_label(doc, "Portfolio at a Glance")
    stats = [
        ("$1.3BN+", "Programme value delivered"),
        ("40+", "Major programmes led"),
        ("13", "Sectors delivered"),
        ("7", "Active ventures (Verax)"),
        ("4", "Proprietary frameworks"),
    ]
    tbl = doc.add_table(rows=2, cols=5)
    tbl.autofit = True
    for i, (big, small) in enumerate(stats):
        c_top = tbl.cell(0, i)
        c_bot = tbl.cell(1, i)
        c_top.text = ""
        c_bot.text = ""
        pt = c_top.paragraphs[0]
        pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(pt, big, bold=True, size=18, color=SIENNA, font="Georgia")
        pb = c_bot.paragraphs[0]
        pb.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(pb, small, size=8.5, color=MUTED)

    # ============================================================
    # PAGE 2 — VALUE & SECTOR COVERAGE
    # ============================================================
    page_break(doc)
    section_label(doc, "How I Add Value")
    pillars = [
        ("Strategic Clarity",
         "Cut through complexity with framework-led thinking — translate vision into a sequenced, "
         "procurement-ready roadmap that boards, ministers and operators can all act on."),
        ("Programme Delivery at Scale",
         "Proven track record running £100M–£1BN+ portfolios across 242 sites and 150+ countries — "
         "directing 450+ cross-functional resources against hard SLA, uptime and ROI commitments."),
        ("Framework-Led Advisory",
         "Author of four proprietary methodologies (Smart Services, Ecosystem, Operations, VRO) "
         "grounded in ISO 37120/37122, BSI PAS 181, ITU-T L.1600 and IEC 62443."),
        ("Innovation Ecosystem Building",
         "Build the connective tissue — demand/supply scoring, funnels, PoC-to-scale governance — "
         "that turns innovation from theatre into compounding, measurable ROI."),
    ]
    for title, text in pillars:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(3)
        add_run(p, f"{title}.  ", bold=True, size=10.5, color=OXFORD)
        add_run(p, text, size=10, color=INK)

    section_label(doc, "Sector Coverage")
    sectors = [
        "Smart Cities — Greenfield", "Smart Cities — Brownfield", "Airports & Aviation",
        "Ports & Maritime", "Industrial / OT & Mining", "Campus & Education",
        "Malls & Retail", "Lifestyle Developments", "Urban Operations",
        "Central Government", "Defence", "Diplomatic / FCO",
    ]
    sec_tbl = doc.add_table(rows=4, cols=3)
    for i, s in enumerate(sectors):
        c = sec_tbl.cell(i // 3, i % 3)
        c.text = ""
        pp = c.paragraphs[0]
        add_run(pp, "▸  ", bold=True, size=10, color=SIENNA)
        add_run(pp, s, size=10, color=INK)

    section_label(doc, "Global Standards & Frameworks Applied")
    body(doc,
         "ISO 37120:2018 (City Services & Quality of Life Indicators)  ·  "
         "ISO 37122:2019 (Smart City Indicators)  ·  "
         "BSI PAS 181 (Smart City Framework for Decision-Makers)  ·  "
         "ITU-T L.1600 / FG-SSC (Smart Sustainable Cities KPI Framework)  ·  "
         "IEC 62443 (Industrial Cybersecurity, OT & ICS)  ·  "
         "ITIL v3  ·  ISO 27001  ·  GDPR  ·  TOGAF ADM  ·  PRINCE2 / MSP",
         size=9.5, space_after=4)

    # ============================================================
    # PAGE 3 — FOUR PROPRIETARY FRAMEWORKS
    # ============================================================
    page_break(doc)
    section_label(doc, "Four Proprietary Frameworks — Verax Venture Studio IP")
    body(doc,
         "Developed inside Verax Venture Studio and distilled from live engagements at DMCC Uptown Dubai, "
         "NEOM Sindalah, Red Sea Global, Strive Services Group and Istanbul Metropolitan Municipality (IBB).",
         size=9.5, space_after=6, color=MUTED)

    frameworks = [
        ("1. Smart Services Advisory", "the HOW of smart-city delivery",
         "3-layer service definition (Sense → Analyse → Act) · 6-stage pathway (Discovery → Identification → "
         "Prioritisation → Design → Procure → Execute) · 7-domain taxonomy (People & Welcome, Smart Energy, "
         "Smart Water, Smart Environment, Robotics, Vision AI & Safety, Smart Facilities) · 21-measure scoring "
         "with 4× Feasibility multiplier → Impact-vs-Feasibility quadrant · 8 delivered outputs per engagement "
         "(Service Catalogue, Architecture Blueprint, Living Lab Plan, Test Framework, TOM, Business Model, "
         "RFP Pack, Governance + KPIs).",
         "Deployed: 150+ smart services mapped and 40+ smart-city initiatives delivered end-to-end — "
         "DMCC Uptown Dubai, NEOM Sindalah, Red Sea Global."),
        ("2. Innovation Ecosystem Advisory", "the WHAT — the pipeline",
         "8-Pillar Ecosystem Model (Vision & Strategic Alignment, R&A, Investment & Funding, Demand Side, "
         "Supply Side, Government Programmes, Talent & Skills, Enablers) · 7-stage Innovation Funnel with "
         "3–4% conversion · dual scoring (Demand 5×26 measures, Supply 5 categories) with Fit ≥80% fast-track · "
         "benchmarks: ~85% implementation success, 100%+ ROI over 3 years, ~14-month payback.",
         "Deployed: 8 Living Labs delivered, 85+ pilots landed — Strive Services Group (KSA) smart-FM demand "
         "pipeline, vendor registry and governance committee established."),
        ("3. Smart City Operations Advisory", "the RUN — steady-state excellence",
         "TOM design, service transition, NOC/SOC architecture and 24/7 operations model (greenfield or "
         "brownfield) · 60+ SOP library (incident, change, escalation, service restoration) mapped to ITIL v3 · "
         "hyper-automation and AI-driven anomaly detection reducing manual incident handling by 40% · "
         "Northbound application layer (citizen apps, operator dashboards, CCC interfaces, OSS/BSS).",
         "Delivered: 99.9% uptime at NEOM  ·  99.8% SLA at FCO across 242 sites / 150+ countries  ·  "
         "Cisco IBB Istanbul architecture blueprint."),
        ("4. Value Realisation Office (VRO)", "the KEEP DELIVERING engine",
         "VRO ≠ PMO — PMO governs delivery (time, scope, cost); VRO governs value (outcomes, ROI, evolution). "
         "3 parallel operating modes (Value Realisation, Operational Excellence, Adaptive/Evolving) · "
         "3 permanent workstreams (Service Portfolio & Business Model, Technology Blueprint, Operating Model) · "
         "3 strategic imperatives · 4 live input feeds · 5 standing outputs (Innovation Stimulus, 12-month "
         "rolling Roadmap, Risk Register, 3-year Investments Forecast, PoC pipeline).",
         "Deployed: NEOM (Innovation governance board), Quantela (outcome-based commercial model), "
         "Cisco global smart city practice."),
    ]
    for num, tag, core, deploy in frameworks:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(1)
        add_run(p, num, bold=True, size=11, color=OXFORD)
        add_run(p, f"  —  {tag}", italic=True, size=10, color=SIENNA)
        body(doc, core, size=9.5, space_after=2)
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(2)
        add_run(p2, "Live deployment.  ", bold=True, size=9.5, color=OXFORD, italic=True)
        add_run(p2, deploy, size=9.5, color=INK, italic=True)

    # Closed-loop
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "Ecosystem (WHAT) → Smart Services (HOW) → Operations (RUN) → VRO (KEEP DELIVERING)",
            bold=True, size=10, color=SIENNA, font="Georgia", italic=True)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(0)
    add_run(p2, "A closed loop — VRO feeds new ideation back into the Ecosystem funnel, "
                "preventing programme drift and sustaining ROI beyond Phase 2.",
            size=9, color=MUTED, italic=True)

    # ============================================================
    # PAGE 4-5 — PROFESSIONAL EXPERIENCE
    # ============================================================
    page_break(doc)
    section_label(doc, "Professional Experience")

    roles = [
        ("Director of Innovation", "NEOM", "Nov 2024 – Present", "£93M", "Riyadh, Saudi Arabia", [
            "Spearhead technology, innovation and partnerships strategy for the world's largest giga-project — "
            "engaging NEOM sectors and regions to identify critical and emerging technology needs across smart "
            "infrastructure, AI, IoT and data.",
            "Govern evaluation of AI, IoT, digital twin, edge and autonomous systems for deployment across "
            "THE LINE, SINDALAH and TROJENA — setting KPIs, SLAs and operational frameworks for pilot programmes.",
            "Cultivate high-impact global partnerships with technology leaders, venture capital firms and "
            "sovereign investors. Awarded 'Best Performing Employee' in a sector of 600+ staff.",
        ]),
        ("Director of Technology & Digital", "NEOM", "Apr 2023 – Nov 2024", "£230M", "Riyadh, Saudi Arabia", [
            "Headhunted to lead 90+ resources delivering NEOM's digital master plan — Smart City platforms, "
            "IoT infrastructure and OT systems across Urban Planning and Islands divisions.",
            "Defined and governed the Cloud-Edge Collaborative infrastructure strategy — multi-cloud "
            "(Azure IoT, AWS), on-premises environments and distributed edge nodes processing real-time IoT data.",
            "Delivered smart building and OT systems (BMS, ACS, CCTV, AV/ELV, structured cabling) end-to-end, "
            "and built the 24/7 NOC capability that achieved 99.9% uptime across critical infrastructure.",
            "Applied Verax Smart Services Advisory at NEOM Sindalah — scored and prioritised services across "
            "7 domains, producing procurement-ready Service Catalogue and Architecture Blueprint.",
        ]),
        ("Founder & Principal — Smart City Advisory", "Verax Venture Studio", "2022 – Present",
         "Concurrent", "UAE / KSA / LATAM", [
            "Founded Verax to codify 20 years of smart city delivery into four vendor-agnostic methodologies: "
            "Smart Services Advisory, Innovation Ecosystem Advisory, Operations Advisory and VRO.",
            "DMCC Uptown Dubai: applied 6-stage pathway to score 40+ candidate services using 21-measure "
            "framework — delivered full procurement-ready portfolio (Blueprint, RFP Pack, Governance & KPIs).",
            "Strive Services Group (KSA): deployed Innovation Ecosystem Advisory to structure smart-FM demand "
            "pipeline — dual demand/supply scoring, vendor registry and governance committee established.",
            "Built a 31-agent hierarchical AI orchestration platform automating email triage, CRM enrichment, "
            "procurement and cost monitoring — eliminating 80%+ of manual operational overhead at $30/month total "
            "platform cost.",
        ]),
        ("Director — Smart Cities & IoT", "Wipro Technologies", "Apr 2022 – Apr 2023", "£80M",
         "Riyadh, Saudi Arabia", [
            "Established Wipro's Arabian entity reporting directly to CEO — strategy, growth and IoT programme "
            "delivery for NEOM and Red Sea Global giga-projects; led 70+ resources.",
            "Delivered KSA's first Smart Luggage Services and advanced EV fleet management — full product "
            "lifecycle from design through to 24/7 NOC-based operational support.",
            "Built operational service delivery capability (NOC, helpdesk, field services) supporting 24/7 OT "
            "infrastructure operations across client sites.",
        ]),
        ("Director — Growth & Delivery", "Quantela (Smart Cities Platform)", "Sept 2021 – Apr 2022", "£98M",
         "Global", [
            "Recruited by Board to drive strategic growth — increased company revenue from $18M to $90M (400%) "
            "with $300M+ pipeline through strategic business model transformation.",
            "Forged key commercial and investor relationships including the $2BN Digital Alpha fund; delivered "
            "first major deal closure through outcome-based, revenue-sharing commercial framework.",
        ]),
        ("Smart City Global Advisory Lead & Programme Director", "Cisco Systems", "Nov 2015 – Apr 2021",
         "£250M", "Global", [
            "Established and scaled Cisco's global Smart Cities and IoT practice from zero — delivering "
            "programmes across 10+ city deployments in Europe, Middle East and Asia-Pacific, individual "
            "programmes up to £45M.",
            "Designed ITIL-aligned operational service models for Smart City platforms — SOPs, incident "
            "management, NOC frameworks and performance optimisation standards.",
            "Provided CXO-level advisory across public sector, utilities and telecoms — translating OT/IT "
            "convergence strategy into scalable platform and connectivity solutions.",
            "Directed product strategy, go-to-market and EMEA systems-integrator partnerships for Cisco's "
            "Connected Cities and IoT portfolio.",
        ]),
        ("Programme Director — Public Service Network", "Cabinet Office", "Mar 2013 – Nov 2015", "£18M p.a.",
         "United Kingdom", [
            "Worked directly with Cabinet Ministers to deliver technology strategy for Public Service "
            "Authorities — managed 48 suppliers and governed Local Authority CTO/CEO/COO leadership.",
            "Established operational service management framework — SLAs, incident, change control and "
            "performance reporting for public sector OT/IT infrastructure.",
        ]),
        ("Programme Director", "Environment Agency", "Dec 2011 – Mar 2013", "£38M", "United Kingdom", [
            "Brought into a failing central government department to redefine enterprise ways of working — "
            "established Enterprise Architecture office and PMO, leading 24 project managers to successful outcomes.",
        ]),
        ("Director, Strategic Operations", "Foreign & Commonwealth Office", "Aug 2008 – Dec 2011", "£768M",
         "Global", [
            "Reported to the Foreign Secretary and CIO — governed global service delivery operations across "
            "242 sites in 150+ countries, single point of contact for all Ambassadors and Deputy Heads of Mission.",
            "Led 24/7 NOC operations, service desk, infrastructure operations and business continuity across "
            "secure diplomatic OT/IT networks — 450+ indirect reports, 99.8% SLA, ISO 27001 compliance.",
        ]),
        ("Programme Director — YJB", "Ministry of Justice", "Feb 2008 – Aug 2008", "£7M", "United Kingdom", [
            "Delivered Youth Justice Board technology programme — governance, commercial control and stakeholder "
            "management across central government and partner agencies.",
        ]),
        ("Overseas Programme Director — DII(F)", "Ministry of Defence", "Nov 2005 – Aug 2008", "£1BN+", "Global", [
            "Led consortium delivery of 5 large enterprise organisations on one of the largest global business "
            "change programmes of its time — Defence Information Infrastructure (Future): 150,000 workstations, "
            "300,000 users, 2,000 sites.",
            "Directed programme management and enterprise architecture across classified defence infrastructure, "
            "working to UK Government TS/SC security standards.",
        ]),
        ("Transformation Lead", "Corus Steel", "Aug 2004 – Nov 2005", "£19M", "United Kingdom", [
            "Led enterprise transformation programme across industrial manufacturing operations — process "
            "re-engineering, systems consolidation and operational readiness for integrated steel sites.",
        ]),
        ("Implementation Lead", "Department for Work and Pensions (DWP)", "Mar 2004 – Aug 2005", "£4M",
         "United Kingdom", [
            "Delivered enterprise implementation across DWP — requirements, design, build and transition into "
            "live public-sector operations.",
        ]),
        ("Deployment Lead", "HM Prison Service", "Jan 2003 – Mar 2004", "£6M", "United Kingdom", [
            "Managed estate-wide deployment programme across HM Prison Service secure sites — rollout planning, "
            "risk management and operational handover under strict security and custodial constraints.",
        ]),
    ]
    for title, org, dates, value, loc, bullets in roles:
        subheading(doc, f"{title}  ·  {org}", f"{dates}  |  {value}  |  {loc}")
        for b in bullets:
            bullet(doc, b, size=9.5)

    # ============================================================
    # PAGE 6 — COMPETENCIES, CREDENTIALS, EDUCATION
    # ============================================================
    page_break(doc)
    section_label(doc, "Core Competencies")
    comp_tbl = doc.add_table(rows=1, cols=2)
    comp_tbl.autofit = True
    c1 = comp_tbl.cell(0, 0)
    c2 = comp_tbl.cell(0, 1)
    c1.text = ""; c2.text = ""

    def comp_col(cell, heading, items):
        p = cell.paragraphs[0]
        add_run(p, heading, bold=True, size=10, color=OXFORD, smallcaps=True, letter_spacing=20)
        for it in items:
            pp = cell.add_paragraph()
            pp.paragraph_format.space_after = Pt(1)
            add_run(pp, "▸  ", bold=True, size=9.5, color=SIENNA)
            add_run(pp, it, size=9.5, color=INK)

    comp_col(c1, "Smart City & Digital Infrastructure", [
        "IoT Architecture & Platform Design (Azure IoT, 5G/NB-IoT)",
        "AI/ML, Vision AI & Big Data Analytics",
        "Smart Energy, Water, Mobility & Facilities",
        "OT/IT Convergence  ·  SCADA / OPC  ·  V2X",
        "Edge & Hybrid Cloud  ·  Multi-Cloud (Azure, AWS)",
        "Urban Digital Twin  ·  BMS / ACS / CCTV / AV-ELV",
        "Enterprise Connectivity  ·  5G / 4G / LTE / NB-IoT",
        "ERP Transformation (SAP, Oracle Fusion, Dynamics 365, AX12)",
        "Blockchain, Web3 & Autonomous Systems",
        "Cybersecurity & Compliance (ISO 27001, IEC 62443, GDPR)",
    ])
    comp_col(c2, "Executive Leadership & Strategy", [
        "C-Suite Advisory & Board-Level Communication",
        "P&L, CAPEX/OPEX & Financial Forecasting",
        "Regulatory & Government Affairs  ·  M&A  ·  PPP",
        "Cross-Functional Team Leadership (450+ resources)",
        "Vendor Negotiation & 48+ Supplier Governance",
        "Innovation Ecosystem Development",
        "Agile  ·  PRINCE2  ·  MSP  ·  ITIL v3  ·  TOGAF ADM",
        "Procurement, RFP Design & Commercial Modelling",
        "Performance Management & KPI Governance",
        "Living Lab / Model Office / TOM Design",
    ])

    section_label(doc, "Certifications")
    body(doc,
         "PRINCE2 Practitioner  ·  MSP (Managing Successful Programmes)  ·  ITIL v3 Service Management  ·  "
         "CompTIA A+ & Security+  ·  BTA Blockchain Professional  ·  BTA Solution Architect  ·  "
         "Microsoft Azure / M365",
         size=10, space_after=4)

    section_label(doc, "Education")
    body(doc, "MBA — Entrepreneurial Action.  Anglia Ruskin University, Cambridge  ·  2025.", space_after=2)
    body(doc, "MBA — Artificial Intelligence (in progress).  University of Cumbria  ·  Expected 2026.",
         space_after=2)
    body(doc, "Postgraduate-Level Strategic Management and Leadership.", space_after=4)

    section_label(doc, "Security & Languages")
    body(doc, "UK Security Clearance: SC / DV eligible (FCO & MoD history).  ·  English (native).  ·  "
              "Saudi Arabia Premium Resident — available for global opportunities.", space_after=4)

    section_label(doc, "Interests")
    body(doc,
         "Technology  ·  Motorbike Racing  ·  Self-Development  ·  Volunteering  ·  CrossFit  ·  Travel & Culture.",
         space_after=4)

    section_label(doc, "References")
    body(doc, "Available on request.", size=9.5, color=MUTED)

    doc.save(DOCX_PATH)
    print(f"DOCX: {DOCX_PATH}  ({DOCX_PATH.stat().st_size:,} bytes)")


def to_pdf():
    """Convert DOCX -> PDF via Word COM (docx2pdf)."""
    try:
        from docx2pdf import convert
        convert(str(DOCX_PATH), str(PDF_PATH))
        print(f"PDF : {PDF_PATH}  ({PDF_PATH.stat().st_size:,} bytes)")
        return True
    except Exception as e:
        print(f"docx2pdf failed: {e}")
        return False


if __name__ == "__main__":
    build()
    to_pdf()
